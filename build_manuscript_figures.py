"""Embed figures 2-7 into the VEG-05 microbiome manuscript and build an illustrated docx.
Inserts figure blocks after anchor sentences and appends a Figure Legends section.
"""
import os, re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from PIL import Image

HERE = os.path.dirname(os.path.abspath(__file__))
md = open(os.path.join(HERE, 'manuscript_draft.md'), encoding='utf-8').read()

# (anchor substring that ends a paragraph)  ->  list of (figure file, caption)
FIGS = [
    ("2.18 versus 0.80 in ground blue-light controls.", [
        ("figures/fig2a_alpha_diversity_16S.png",
         "**Figure 2A.** Bacterial (16S) alpha diversity (Observed ASVs) by flight status and light treatment; flight leaves show higher diversity.")]),
    ("Adventitious root fungal communities were significantly reshaped (R²=0.65, F=8.19, p=0.001).", [
        ("figures/fig2b_dysbiosis_16S.png",
         "**Figure 2B (16S).** Bacterial dysbiosis index (Bray-Curtis displacement from ground centroid); elevated under flight."),
        ("figures/fig2b_dysbiosis_ITS.png",
         "**Figure 2B (ITS).** Fungal dysbiosis index; the largest displacement is in flight leaves.")]),
    ("the root transcriptional response to spaceflight is less light-dependent than the leaf response.", [
        ("figures/fig3_deg_summary.png",
         "**Figure 3.** DEG summary across tissue × contrast; the leaf blue-light flight response (4,716 DEGs) dwarfs red (523)."),
        ("figures/fig3b_volcano_leaf_flight.png", "**Figure 3B.** Leaf Flight-vs-Ground volcano."),
        ("figures/fig3c_volcano_advroot_flight.png", "**Figure 3C.** Adventitious-root Flight-vs-Ground volcano (predominantly upregulation).")]),
    ("consistent with spaceflight-induced oxidative stress.", [
        ("figures/fig4_module_traits_Leaf.png",
         "**Figure 4 (Leaf).** WGCNA module–trait correlations (flight, light, 16S/ITS dysbiosis)."),
        ("figures/fig4_module_traits_AdvRoot.png",
         "**Figure 4 (Adv-Root).** Module–trait correlations; the black module tracks ITS dysbiosis (r=−0.85) but not flight.")]),
    ("microbial features were consistently present in the top weights.", [
        ("figures/fig5a_mofa_variance.png", "**Figure 5A.** MOFA+ variance explained per factor per view (transcriptome, 16S, ITS)."),
        ("figures/fig5b_mofa_correlations.png", "**Figure 5B.** MOFA+ factor–trait correlations; Factor 1 captures flight (ρ=−0.76).")]),
    ("*Paenibacillus* (ρ=−0.77, padj=0.003).", [
        ("figures/fig6_module_taxon_network.png",
         "**Figure 6.** Bipartite module–taxon network linking flight-associated modules to Methylobacterium, Burkholderia, Azospirillum.")]),
    ("as the main ecological guilds.", [
        ("figures/fig7_faprotax.png",
         "**Figure 7.** FAPROTAX predicted bacterial functions (aerobic chemoheterotrophy, N fixation, methanotrophy, methanol oxidation).")]),
]

def fig_block(items):
    out = "\n"
    for path, cap in items:
        out += f"\n![{cap.split('.**')[0].replace('**','')}]({path})\n\n*{cap}*\n"
    return out

inserted = 0
for anchor, items in FIGS:
    if anchor in md:
        md = md.replace(anchor, anchor + fig_block(items), 1)
        inserted += len(items)
    else:
        print("WARN anchor not found:", anchor[:50])

legends = "\n---\n\n## Figure legends\n\n"
for _, items in FIGS:
    for _, cap in items:
        legends += "- " + cap.replace("**", "**") + "\n"
md = md.replace("## Acknowledgements", legends + "\n---\n\n## Acknowledgements", 1)

open(os.path.join(HERE, 'manuscript_with_figures.md'), 'w', encoding='utf-8').write(md)
print("figures embedded:", inserted)

# ---- build docx ----
doc = Document(); s = doc.sections[0]
s.page_width, s.page_height = Inches(8.5), Inches(11)
for m in ('top', 'bottom', 'left', 'right'):
    setattr(s, f'{m}_margin', Inches(1))
doc.styles['Normal'].font.name = 'Arial'; doc.styles['Normal'].font.size = Pt(11)
INLINE = re.compile(r'(\*\*.+?\*\*|\*.+?\*)')
IMG = re.compile(r'^!\[.*?\]\((.+?)\)\s*$')

def add_runs(par, text):
    for tok in INLINE.split(text):
        if not tok:
            continue
        if tok.startswith('**') and tok.endswith('**'):
            par.add_run(tok[2:-2]).bold = True
        elif tok.startswith('*') and tok.endswith('*'):
            par.add_run(tok[1:-1]).italic = True
        else:
            par.add_run(tok)

for line in md.splitlines():
    t = line.rstrip()
    if not t.strip():
        continue
    mi = IMG.match(t)
    if mi:
        p = os.path.normpath(os.path.join(HERE, mi.group(1)))
        if os.path.exists(p):
            w, h = Image.open(p).size
            width = Inches(6.3 if h / w <= 1.15 else min(6.3, 8.2 * w / h))
            par = doc.add_paragraph(); par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            par.add_run().add_picture(p, width=width)
        continue
    if t.startswith('# '):
        add_runs(doc.add_heading(level=0), t[2:])
    elif t.startswith('### '):
        add_runs(doc.add_heading(level=2), t[4:])
    elif t.startswith('## '):
        add_runs(doc.add_heading(level=1), t[3:])
    elif t == '---':
        continue
    elif t.startswith('- '):
        add_runs(doc.add_paragraph(style='List Bullet'), t[2:])
    else:
        par = doc.add_paragraph()
        if t.startswith('*') and t.endswith('*') and not t.startswith('**'):
            r = par.add_run(t.strip('*')); r.italic = True; r.font.size = Pt(9.5)
        else:
            add_runs(par, t)

z = doc.settings.element.find(qn('w:zoom'))
if z is not None and z.get(qn('w:percent')) is None:
    z.set(qn('w:percent'), '100')
out = os.path.join(HERE, 'manuscript_with_figures.docx')
doc.save(out)
print("docx:", os.path.getsize(out), "bytes, images:", sum(1 for _ in doc.inline_shapes))
